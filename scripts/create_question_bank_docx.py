import argparse
import json
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_BANK_PATH = ROOT / "data" / "question-bank.json"
DEFAULT_OUT_PATH = ROOT / "question_bank_review.docx"

FONT_BODY = "Microsoft YaHei"
COLOR_INK = RGBColor(35, 39, 44)
COLOR_MUTED = RGBColor(102, 112, 125)
COLOR_BLUE = RGBColor(31, 91, 142)
COLOR_GREEN = RGBColor(39, 124, 86)


def cjk_count(text):
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def repair_text(value):
    text = "" if value is None else str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return text
    if not any(ord(ch) > 127 for ch in text):
        return text
    if cjk_count(text) > 0:
        return text
    try:
        repaired = text.encode("latin1").decode("utf-8")
    except UnicodeError:
        return text
    if cjk_count(repaired) >= max(1, cjk_count(text)):
        return repaired
    return text


def set_run_font(run, name=FONT_BODY, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_run(paragraph, text, size=10.5, bold=False, color=COLOR_INK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def set_shading(target, fill):
    if hasattr(target, "_p"):
        props = target._p.get_or_add_pPr()
    else:
        props = target._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    props.append(header)


def set_cell_width(cell, cm):
    props = cell._tc.get_or_add_tcPr()
    width = props.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        props.append(width)
    width.set(qn("w:w"), str(int(cm * 567)))
    width.set(qn("w:type"), "dxa")


def question_type(question):
    raw = repair_text(question.get("type") or "")
    if raw in {"MultipleChoice", "2"} or "Multiple" in raw:
        return "Multiple Choice"
    if raw in {"Judgement", "Judge", "6"} or "Judgement" in raw or "Judg" in raw:
        return "True / False"
    if raw in {"SingleChoice", "1"} or "Single" in raw:
        return "Single Choice"
    return raw or "Unclassified"


def clean_options(question):
    options = []
    for index, option in enumerate(question.get("options") or []):
        label = repair_text(option.get("label") or chr(ord("A") + index)).upper()
        text = repair_text(option.get("text") or "")
        if text:
            options.append({"label": label, "text": text})
    return options


def answer_labels(question, options):
    labels = {repair_text(label).upper() for label in question.get("correctLabels") or [] if repair_text(label)}
    if labels:
        return sorted(labels)

    by_text = {option["text"].replace(" ", ""): option["label"] for option in options}
    resolved = set()
    for raw in question.get("correctTexts") or []:
        text = repair_text(raw)
        compact = text.replace(" ", "")
        if compact in by_text:
            resolved.add(by_text[compact])
        elif compact.lower() == "true":
            for option in options:
                if any(word in option["text"] for word in ("正确", "对", "是")):
                    resolved.add(option["label"])
        elif compact.lower() == "false":
            for option in options:
                if any(word in option["text"] for word in ("错误", "错", "否")):
                    resolved.add(option["label"])
    return sorted(resolved)


def answer_short(question):
    options = clean_options(question)
    labels = answer_labels(question, options)
    if labels:
        return ", ".join(labels)
    texts = [repair_text(text) for text in question.get("correctTexts") or [] if repair_text(text)]
    return ", ".join(texts) or "Not available"


def answer_full(question):
    options = clean_options(question)
    option_by_label = {option["label"]: option for option in options}
    pieces = []
    for label in answer_labels(question, options):
        option = option_by_label.get(label)
        pieces.append(f"{label}. {option['text']}" if option else label)
    if pieces:
        return "; ".join(pieces)
    return "; ".join(repair_text(text) for text in question.get("correctTexts") or []) or "Not available"


def add_heading_band(doc, title, subtitle=""):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = True
    set_shading(paragraph, "DDEBF7")
    add_run(paragraph, f"  {title}", size=15, bold=True, color=COLOR_BLUE)
    if subtitle:
        add_run(paragraph, f"  {subtitle}", size=9, color=COLOR_MUTED)


def add_quick_index(doc, questions):
    add_heading_band(doc, "Answer Index", "Original ordering from the bank")
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    for index, cell in enumerate(table.rows[0].cells):
        set_shading(cell, "1F5B8E")
        set_cell_margins(cell, top=80, bottom=80, start=70, end=70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(paragraph, "No." if index % 2 == 0 else "Answer", size=9, bold=True, color=RGBColor(255, 255, 255))

    for row_start in range(0, len(questions), 4):
        row = table.add_row()
        for block in range(4):
            question_index = row_start + block
            number_cell = row.cells[block * 2]
            answer_cell = row.cells[block * 2 + 1]
            for cell in (number_cell, answer_cell):
                set_cell_margins(cell, top=65, bottom=65, start=70, end=70)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            number_cell.text = ""
            number_p = number_cell.paragraphs[0]
            number_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if question_index < len(questions):
                add_run(number_p, str(question_index + 1), size=9, bold=True)
                answer_cell.text = ""
                answer_p = answer_cell.paragraphs[0]
                answer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(answer_p, answer_short(questions[question_index]), size=8.5, color=COLOR_BLUE)
            else:
                answer_cell.text = ""

    widths = [0.9, 1.3] * 4
    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)


def add_question_block(doc, index, question):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    add_run(p, f"{index}. ", size=11, bold=True, color=COLOR_BLUE)
    add_run(p, repair_text(question.get("stem") or ""), size=11, bold=True)
    add_run(p, f"  [{question_type(question)}]", size=9, color=COLOR_MUTED)

    options = clean_options(question)
    for option in options:
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Cm(0.4)
        op.paragraph_format.space_before = Pt(0)
        op.paragraph_format.space_after = Pt(0)
        label = option["label"]
        text = option["text"]
        add_run(op, f"{label}. ", size=10.2, bold=True, color=COLOR_GREEN)
        add_run(op, text, size=10.2)

    ans = doc.add_paragraph()
    ans.paragraph_format.left_indent = Cm(0.4)
    ans.paragraph_format.space_before = Pt(1)
    ans.paragraph_format.space_after = Pt(2)
    add_run(ans, "Answer: ", size=10, bold=True, color=COLOR_BLUE)
    add_run(ans, answer_full(question), size=10, color=COLOR_INK)

    if question.get("explanation"):
        exp = doc.add_paragraph()
        exp.paragraph_format.left_indent = Cm(0.4)
        exp.paragraph_format.space_before = Pt(1)
        exp.paragraph_format.space_after = Pt(5)
        add_run(exp, "Note: ", size=9.5, bold=True, color=COLOR_MUTED)
        add_run(exp, repair_text(question.get("explanation")), size=9.5, color=COLOR_MUTED)


def load_bank(path):
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    if isinstance(data, dict) and isinstance(data.get("questions"), dict):
        questions = list(data["questions"].values())
    elif isinstance(data, list):
        questions = data
    else:
        raise ValueError("Unsupported bank format")
    questions = [q for q in questions if isinstance(q, dict)]
    questions.sort(key=lambda q: (question_type(q), repair_text(q.get("stem") or "")))
    return questions


def build_doc(questions, title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_BODY
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    styles["Normal"].font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    add_run(title_p, title, size=18, bold=True, color=COLOR_INK)

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(6)
        add_run(sub_p, subtitle, size=9.5, color=COLOR_MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    counts = Counter(question_type(q) for q in questions)
    add_run(
        meta,
        f"Total {len(questions)} questions | Single {counts.get('Single Choice', 0)} | Multiple {counts.get('Multiple Choice', 0)} | True/False {counts.get('True / False', 0)} | Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=9,
        color=COLOR_MUTED,
    )

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    add_run(
        intro,
        "This file is for local review only. It is generated from the local question bank and should not be published together with any private account data, cookies, screenshots, raw responses, or course URLs.",
        size=9.2,
        color=COLOR_MUTED,
    )

    add_quick_index(doc, questions)

    for group_name in ("Single Choice", "Multiple Choice", "True / False", "Unclassified"):
        group = [q for q in questions if question_type(q) == group_name]
        if not group:
            continue
        add_heading_band(doc, group_name, f"{len(group)} items")
        for index, question in enumerate(group, 1):
            add_question_block(doc, index, question)

    return doc


def main():
    parser = argparse.ArgumentParser(description="Create a clean DOCX review file from a local question bank.")
    parser.add_argument("--bank", default=str(DEFAULT_BANK_PATH), help="Path to the local question-bank JSON file.")
    parser.add_argument("--out", default=str(DEFAULT_OUT_PATH), help="Output DOCX path.")
    parser.add_argument("--title", default="Practice Question Bank Review", help="Document title.")
    parser.add_argument("--subtitle", default="Local review only", help="Document subtitle.")
    args = parser.parse_args()

    questions = load_bank(args.bank)
    doc = build_doc(questions, args.title, args.subtitle)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"saved {out_path} ({len(questions)} questions)")


if __name__ == "__main__":
    main()
